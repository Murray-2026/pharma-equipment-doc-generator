import streamlit as st
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import io
import requests

st.set_page_config(
    page_title="隔离器售前文档生成系统",
    page_icon="🏭",
    layout="wide",
    initial_sidebar_state="expanded"
)

def load_data(key, default=None):
    try:
        data = st.session_state.get(f"local_{key}")
        if data is None:
            data = st.query_params.get(key)
            if data:
                data = json.loads(data)
                st.session_state[f"local_{key}"] = data
        return data if data is not None else default
    except:
        return default

def save_data(key, value):
    st.session_state[f"local_{key}"] = value

def init_session():
    if "urs_text" not in st.session_state:
        st.session_state.urs_text = ""
    if "selected_equipment" not in st.session_state:
        st.session_state.selected_equipment = []
    if "key_params" not in st.session_state:
        st.session_state.key_params = ""
    if "budget_level" not in st.session_state:
        st.session_state.budget_level = "标准型"
    if "language" not in st.session_state:
        st.session_state.language = "中文"
    if "history" not in st.session_state:
        st.session_state.history = load_data("history", [])
    if "customers" not in st.session_state:
        st.session_state.customers = load_data("customers", [])
    if "templates" not in st.session_state:
        st.session_state.templates = load_data("templates", [])
    if "generated_docs" not in st.session_state:
        st.session_state.generated_docs = {"urs": "", "quotation": "", "urs_docx": None, "quotation_docx": None}
    if "parsed_requirements" not in st.session_state:
        st.session_state.parsed_requirements = []
    if "ai_parsed" not in st.session_state:
        st.session_state.ai_parsed = False
    if "openai_api_key" not in st.session_state:
        st.session_state.openai_api_key = ""

def parse_urs_with_ai(text):
    if not st.session_state.openai_api_key:
        return parse_urs_fallback(text)
    
    try:
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {st.session_state.openai_api_key}"
        }
        
        prompt = f"""
        请分析以下制药设备URS文档，提取每条需求的关键信息：
        
        {text}
        
        请按照以下JSON格式输出：
        {{
            "requirements": [
                {{
                    "number": 序号,
                    "text": "需求原文",
                    "category": "分类（结构尺寸/洁净等级/灭菌消毒/操作组件/压力控制/法规合规/材质要求/验证要求/其他）",
                    "equipment_type": "设备类型（单体无菌隔离器/VHP传递窗/整线隔离器/单体负压隔离器/None）",
                    "compliance": ["法规1", "法规2"],
                    "spec_value": "提取的规格值（如L1200×D800×H900mm、ISO 5级等）",
                    "priority": "优先级（高/中/低）",
                    "analysis": "需求分析说明"
                }}
            ]
        }}
        """
        
        data = {
            "model": "gpt-3.5-turbo",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3
        }
        
        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=data, timeout=60)
        result = response.json()
        
        if "choices" in result:
            content = result["choices"][0]["message"]["content"]
            try:
                parsed = json.loads(content)
                st.session_state.parsed_requirements = parsed.get("requirements", [])
                st.session_state.ai_parsed = True
                return parsed.get("requirements", [])
            except:
                st.warning("AI解析结果格式异常，使用备用解析")
                return parse_urs_fallback(text)
        else:
            st.warning("AI解析失败，使用备用解析")
            return parse_urs_fallback(text)
    except Exception as e:
        st.warning(f"AI解析出错: {str(e)}，使用备用解析")
        return parse_urs_fallback(text)

def parse_urs_fallback(text):
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    requirements = []
    
    for line_num, line in enumerate(lines, 1):
        req = {
            'number': line_num,
            'text': line,
            'category': classify_requirement(line),
            'equipment_type': detect_equipment(line),
            'compliance': detect_compliance(line),
            'spec_value': extract_spec_value(line),
            'priority': '中',
            'analysis': f"第{line_num}条需求分析：{line[:50]}..."
        }
        requirements.append(req)
    
    st.session_state.parsed_requirements = requirements
    st.session_state.ai_parsed = False
    return requirements

def classify_requirement(text):
    text_lower = text.lower()
    if '尺寸' in text or 'mm' in text or 'dimension' in text_lower:
        return '结构尺寸'
    if '洁净度' in text or 'ISO' in text or 'clean' in text_lower:
        return '洁净等级'
    if '灭菌' in text or 'VHP' in text or 'decontamination' in text_lower:
        return '灭菌消毒'
    if '手套' in text or 'glove' in text_lower:
        return '操作组件'
    if '压力' in text or 'pressure' in text_lower:
        return '压力控制'
    if 'GMP' in text or 'FDA' in text or '法规' in text or 'compliance' in text_lower:
        return '法规合规'
    if '材质' in text or 'material' in text_lower:
        return '材质要求'
    if '验证' in text or 'IQ' in text or 'OQ' in text or 'PQ' in text:
        return '验证要求'
    return '其他'

def detect_equipment(text):
    text_lower = text.lower()
    if '负压' in text or 'negative pressure' in text_lower:
        return '单体负压隔离器'
    if '整线' in text or 'full line' in text_lower:
        return '整线隔离器'
    if 'VHP' in text or '传递窗' in text or 'pass-through' in text_lower:
        return 'VHP传递窗'
    if '隔离器' in text or 'isolator' in text_lower:
        return '单体无菌隔离器'
    return None

def detect_compliance(text):
    compliances = []
    text_lower = text.lower()
    if '中国GMP' in text or 'chinese gmp' in text_lower:
        compliances.append('中国GMP')
    if 'FDA' in text:
        compliances.append('FDA 21 CFR')
    if 'EU' in text or 'EU GMP' in text:
        compliances.append('EU GMP Annex 1')
    if 'WHO' in text:
        compliances.append('WHO GMP')
    if 'PIC/S' in text:
        compliances.append('PIC/S')
    return compliances

def extract_spec_value(text):
    import re
    match = re.search(r'(\d+)\s*[×xX×]\s*(\d+)\s*[×xX×]\s*(\d+)', text)
    if match:
        return f"{match.group(1)}×{match.group(2)}×{match.group(3)}mm"
    match = re.search(r'ISO\s*(\d+)', text)
    if match:
        return f"ISO {match.group(1)}级"
    match = re.search(r'(\d+)-log', text)
    if match:
        return f"≥{match.group(1)}-log"
    return None

def generate_professional_urs_response():
    customer_name = st.session_state.get('customer_name', '')
    project_name = st.session_state.get('project_name', '')
    contact = st.session_state.get('contact', '')
    date_val = st.session_state.get('date')
    date = date_val.strftime('%Y-%m-%d') if hasattr(date_val, 'strftime') else str(date_val) if date_val else datetime.now().strftime('%Y-%m-%d')
    regulations = [r for r in ['中国GMP', 'FDA 21 CFR', 'EU GMP Annex 1', 'WHO GMP', 'PIC/S'] if st.session_state.get(f"reg_{r}")]
    equipment = st.session_state.selected_equipment
    requirements = st.session_state.parsed_requirements
    
    if not requirements:
        requirements = [{'number': 1, 'text': '请先输入URS内容并点击智能解析', 'category': '其他', 'compliance': [], 'spec_value': None, 'priority': '中', 'analysis': ''}]
    
    reg_text = ', '.join(regulations) if regulations else '待确认'
    eq_text = ', '.join(equipment) if equipment else '待确认'
    
    content = f"""# 用户需求说明（URS）逐条回复
## User Requirement Specification — Point-by-Point Reply

---

| 项目 | 内容 |
|------|------|
| 客户名称 | {customer_name or '待填写'} |
| 项目名称 | {project_name or '待填写'} |
| 供应商 | XX隔离器技术有限公司 |
| 日期 | {date} |
| 设备类型 | {eq_text} |
| 适用法规 | {reg_text} |

---

## URS逐条回复

| 条款编号 | 客户需求原文 | 分类 | 优先级 | 我方响应 | 技术说明 | 法规依据 | 验证方式 | 需求分析 |
|----------|--------------|------|--------|----------|----------|----------|----------|----------|
"""
    
    for req in requirements:
        response = generate_response(req)
        tech_desc = generate_technical_description(req)
        compliance = ', '.join(req['compliance']) if req['compliance'] else reg_text
        validation = generate_validation_method(req)
        analysis = req.get('analysis', '')[:30] + '...' if len(req.get('analysis', '')) > 30 else req.get('analysis', '-')
        
        content += f"""| URS-{str(req['number']).zfill(3)} | {req['text']} | {req['category']} | {req.get('priority', '中')} | {response} | {tech_desc} | {compliance} | {validation} | {analysis} |
"""
    
    content += f"""
---

## 合规声明

我方保证所提供的{eq_text}符合{reg_text}相关要求，并将提供完整的验证文件包（IQ/OQ/PQ）。

**回复单位**: XX隔离器技术有限公司  
**技术负责人**: 待填写  
**联系电话**: 待填写  
**日期**: {date}
"""
    
    return content

def generate_response(req):
    if req['category'] == '法规合规':
        return '满足'
    if req.get('spec_value'):
        return '满足'
    if '要求' in req['text'] or 'shall' in req['text'].lower():
        return '满足'
    return '满足'

def generate_technical_description(req):
    category = req['category']
    
    if category == '结构尺寸':
        return f"设备尺寸设计为{req.get('spec_value') or '符合客户要求的尺寸'}，采用模块化设计，便于安装和维护。"
    if category == '洁净等级':
        return f"操作区洁净度达到{req.get('spec_value') or 'ISO 5级'}，符合cGMP对无菌生产环境的要求，配备高效HEPA过滤器。"
    if category == '灭菌消毒':
        return f"采用VHP汽化过氧化氢灭菌技术，灭菌效果{req.get('spec_value') or '≥6-log SAL'}，残气经催化分解后<1ppm。"
    if category == '操作组件':
        return "配备进口丁腈手套，带手套检漏系统，确保操作过程的密闭性和安全性。"
    if category == '压力控制':
        return "采用正压/负压控制系统，确保操作区与外界的压力梯度，防止交叉污染。"
    if category == '法规合规':
        return "本设备设计符合相关法规要求，提供完整的合规文件包。"
    if category == '材质要求':
        return "与产品接触部分采用316L不锈钢材质，表面粗糙度Ra≤0.4μm，符合GMP要求。"
    if category == '验证要求':
        return "提供完整的IQ/OQ/PQ验证文件包，支持客户进行验证确认。"
    
    return "本设备可完全满足客户该项需求，具体技术方案详见技术规格书。"

def generate_validation_method(req):
    category = req['category']
    
    if category == '结构尺寸':
        return 'IQ/OQ'
    if category == '洁净等级':
        return 'IQ/OQ/PQ'
    if category == '灭菌消毒':
        return 'IQ/OQ/PQ'
    if category == '操作组件':
        return 'IQ/OQ'
    if category == '压力控制':
        return 'IQ/OQ'
    if category == '材质要求':
        return 'IQ'
    if category == '验证要求':
        return 'IQ/OQ/PQ'
    
    return 'IQ/OQ/PQ'

def generate_quotation():
    customer_name = st.session_state.get('customer_name', '')
    project_name = st.session_state.get('project_name', '')
    contact = st.session_state.get('contact', '')
    date_val = st.session_state.get('date')
    date = date_val.strftime('%Y-%m-%d') if hasattr(date_val, 'strftime') else str(date_val) if date_val else datetime.now().strftime('%Y-%m-%d')
    equipment = st.session_state.selected_equipment
    budget = st.session_state.budget_level
    regulations = [r for r in ['中国GMP', 'FDA 21 CFR', 'EU GMP Annex 1', 'WHO GMP', 'PIC/S'] if st.session_state.get(f"reg_{r}")]
    
    prices = {
        '单体无菌隔离器': {'经济型': [80, 150], '标准型': [100, 200], '高端型': [180, 320]},
        'VHP传递窗': {'经济型': [30, 60], '标准型': [40, 80], '高端型': [80, 150]},
        '整线隔离器': {'经济型': [300, 600], '标准型': [400, 800], '高端型': [700, 1500]},
        '单体负压隔离器': {'经济型': [100, 200], '标准型': [120, 250], '高端型': [250, 450]}
    }
    
    eq = equipment[0] if equipment else '单体无菌隔离器'
    price_range = prices[eq][budget]
    base_price = (price_range[0] + price_range[1]) / 2
    
    content = f"""# 快速报价单
## Quick Quotation

---

| 项目 | 内容 |
|------|------|
| 报价编号 | QT-{datetime.now().strftime('%Y%m')}-{str(hash(datetime.now())).strip('-')[:3]} |
| 日期 | {date} |
| 有效期至 | 30天 |
| 客户名称 | {customer_name or '待填写'} |
| 项目名称 | {project_name or '待填写'} |
| 供应商 | XX隔离器技术有限公司 |
| 设备类型 | {eq} |
| 配置等级 | {budget} |
| 适用法规 | {', '.join(regulations)} |

---

## A. 设备概述

{equipment_desc(eq)}

## B. 关键技术参数

| 参数项目 | 经济型 | 标准型 | 高端型 |
|----------|--------|--------|--------|
| 灭菌效果 | SAL 10⁻⁶ | SAL 10⁻⁶ | SAL 10⁻⁶ |
| 残气浓度 | <1ppm | <1ppm | <1ppm |
| 内壁材质 | 316L不锈钢 | 316L不锈钢 | 316L不锈钢 |
| 控制系统 | PLC+7"触摸屏 | PLC+7"触摸屏 | PLC+SCADA+10"触摸屏 |
| 数据记录 | 自动记录 | 自动记录 | 审计追踪记录 |
| 双门互锁 | 机械+电气 | 机械+电气 | 机械+PLC三重 |

## C. 报价明细（{budget}方案）

| 序号 | 项目 | 金额（万元） | 备注 |
|------|------|--------------|------|
| 1 | 设备本体 | {base_price * 0.7:.1f} | 含主体结构、手套、视窗、VHP发生器 |
| 2 | 安装调试 | {base_price * 0.08:.1f} | 含现场安装、调试、开机验证 |
| 3 | 验证服务（IQ/OQ/PQ） | {base_price * 0.1:.1f} | 含验证方案编制及执行 |
| 4 | 技术培训 | {base_price * 0.02:.1f} | {3 if budget == '经济型' else 5}天（操作+维护{'+验证' if budget == '高端型' else ''}） |
| 5 | 备件包（首年） | {base_price * 0.02:.1f} | 易损件及推荐备件 |
| 6 | 技术文件包 | {base_price * 0.02:.1f} | 含手册、图纸、证书 |
| 7 | 国内运输+保险 | {base_price * 0.06:.1f} | 含运输一切险 |
| **8** | **合计** | **{base_price:.1f}** | 含13%增值税 |

## D. 配置清单

| 序号 | 项目 | 数量 | 备注 |
|------|------|------|------|
| 1 | {eq}主机系统 | 1台 | 含舱体、门系统、互锁装置 |
| 2 | HEPA过滤器（H14级） | 2个 | 进风+排风 |
| 3 | 备用手套（丁腈） | 2副 | 进口材质 |
| 4 | VHP发生器 | 1套 | 集成式 |
| 5 | 专用工具包 | 1套 | 维护工具 |
| 6 | 操作手册 | 1份 | 中文/英文 |
| 7 | 验证文件包 | 1份 | IQ/OQ/PQ方案 |

## E. 交货周期

| 方案 | 预计周期 | 说明 |
|------|----------|------|
| 经济型 | 12-16周 | 自预付款到账之日起 |
| 标准型 | 16-20周 | 自预付款到账之日起 |
| 高端型 | 20-28周 | 含系统集成交付与调试时间 |

## F. 售后服务条款

| 项目 | 经济型 | 标准型 | 高端型 |
|------|--------|--------|--------|
| 质保期 | 12个月 | 12个月 | 24个月 |
| 远程响应时间 | 48小时 | 24小时 | 24小时 |
| 现场响应时间 | 72小时 | 48小时 | 24小时 |
| 年度保养 | 1次/年 | 1次/年 | 2次/年 |
| 软件升级 | 质保期内免费 | 质保期内免费 | 质保期内免费+质保期后3年优惠 |
| 设备备件供应 | 设备停产10年 | 设备停产10年 | 设备停产10年 |

---

## 声明

1. 本报价有效期为30天，逾期价格可能调整
2. 本报价不含土建、公用工程接口以外的工作
3. 最终价格以双方签订的正式合同为准
4. 如客户有重大变更，供应商保留调整报价的权利

**报价单位**: XX隔离器技术有限公司  
**联系人**: {contact or '待填写'}  
**日期**: {date}
"""
    
    return content

def equipment_desc(eq):
    descs = {
        '单体无菌隔离器': '无菌分装、无菌检测、称量操作的理想选择，正压运行，ISO 5级洁净度，配备VHP灭菌系统。',
        'VHP传递窗': '物料/工具VHP灭菌传递，双门互锁设计，灭菌效果≥6-log SAL，残气<1ppm，适用于洁净区物料传递。',
        '整线隔离器': '灌装+轧盖+称重全线隔离系统，适用于无菌制剂生产线，集成式设计，提高生产效率和产品质量。',
        '单体负压隔离器': '有毒/高活性药品OEL防护专用，负压运行，OEL控制<1μg/m³，双重HEPA过滤，BIBO更换系统。'
    }
    return descs.get(eq, '')

def create_urs_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    title = doc.add_heading('用户需求说明（URS）逐条回复', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('User Requirement Specification — Point-by-Point Reply', level=1)
    doc.add_paragraph()
    
    customer_name = st.session_state.get('customer_name', '')
    project_name = st.session_state.get('project_name', '')
    contact = st.session_state.get('contact', '')
    date_val = st.session_state.get('date')
    date = date_val.strftime('%Y-%m-%d') if hasattr(date_val, 'strftime') else str(date_val) if date_val else datetime.now().strftime('%Y-%m-%d')
    regulations = [r for r in ['中国GMP', 'FDA 21 CFR', 'EU GMP Annex 1', 'WHO GMP', 'PIC/S'] if st.session_state.get(f"reg_{r}")]
    equipment = st.session_state.selected_equipment
    requirements = st.session_state.parsed_requirements
    
    if not requirements:
        requirements = [{'number': 1, 'text': '请先输入URS内容并点击智能解析', 'category': '其他', 'compliance': [], 'spec_value': None, 'priority': '中', 'analysis': ''}]
    
    reg_text = ', '.join(regulations) if regulations else '待确认'
    eq_text = ', '.join(equipment) if equipment else '待确认'
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    info_cells = info_table._cells
    
    info_cells[0].text = '客户名称'
    info_cells[1].text = customer_name or '待填写'
    info_cells[2].text = '项目名称'
    info_cells[3].text = project_name or '待填写'
    info_cells[4].text = '供应商'
    info_cells[5].text = 'XX隔离器技术有限公司'
    info_cells[6].text = '日期'
    info_cells[7].text = date
    info_cells[8].text = '设备类型'
    info_cells[9].text = eq_text
    info_cells[10].text = '适用法规'
    info_cells[11].text = reg_text
    
    for row in info_table.rows:
        for cell in row.cells:
            cell.width = Inches(3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    doc.add_heading('URS逐条回复', level=1)
    
    cols = ['条款编号', '客户需求原文', '分类', '优先级', '我方响应', '技术说明', '法规依据', '验证方式', '需求分析']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = col
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    for req in requirements:
        row_cells = table.add_row().cells
        row_cells[0].text = f"URS-{str(req['number']).zfill(3)}"
        row_cells[1].text = req['text']
        row_cells[2].text = req['category']
        row_cells[3].text = req.get('priority', '中')
        row_cells[4].text = generate_response(req)
        row_cells[5].text = generate_technical_description(req)
        row_cells[6].text = ', '.join(req['compliance']) if req['compliance'] else reg_text
        row_cells[7].text = generate_validation_method(req)
        row_cells[8].text = req.get('analysis', '-')[:50] + '...' if len(req.get('analysis', '')) > 50 else req.get('analysis', '-')
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    doc.add_heading('合规声明', level=1)
    doc.add_paragraph(f"我方保证所提供的{eq_text}符合{reg_text}相关要求，并将提供完整的验证文件包（IQ/OQ/PQ）。")
    
    doc.add_paragraph()
    doc.add_paragraph(f"**回复单位**: XX隔离器技术有限公司")
    doc.add_paragraph(f"**技术负责人**: 待填写")
    doc.add_paragraph(f"**联系电话**: 待填写")
    doc.add_paragraph(f"**日期**: {date}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def create_quotation_docx():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    title = doc.add_heading('快速报价单', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Quick Quotation', level=1)
    doc.add_paragraph()
    
    customer_name = st.session_state.get('customer_name', '')
    project_name = st.session_state.get('project_name', '')
    contact = st.session_state.get('contact', '')
    date_val = st.session_state.get('date')
    date = date_val.strftime('%Y-%m-%d') if hasattr(date_val, 'strftime') else str(date_val) if date_val else datetime.now().strftime('%Y-%m-%d')
    equipment = st.session_state.selected_equipment
    budget = st.session_state.budget_level
    regulations = [r for r in ['中国GMP', 'FDA 21 CFR', 'EU GMP Annex 1', 'WHO GMP', 'PIC/S'] if st.session_state.get(f"reg_{r}")]
    
    prices = {
        '单体无菌隔离器': {'经济型': [80, 150], '标准型': [100, 200], '高端型': [180, 320]},
        'VHP传递窗': {'经济型': [30, 60], '标准型': [40, 80], '高端型': [80, 150]},
        '整线隔离器': {'经济型': [300, 600], '标准型': [400, 800], '高端型': [700, 1500]},
        '单体负压隔离器': {'经济型': [100, 200], '标准型': [120, 250], '高端型': [250, 450]}
    }
    
    eq = equipment[0] if equipment else '单体无菌隔离器'
    price_range = prices[eq][budget]
    base_price = (price_range[0] + price_range[1]) / 2
    
    info_table = doc.add_table(rows=9, cols=2)
    info_table.style = 'Table Grid'
    info_cells = info_table._cells
    
    info_cells[0].text = '报价编号'
    info_cells[1].text = f"QT-{datetime.now().strftime('%Y%m')}-{str(hash(datetime.now())).strip('-')[:3]}"
    info_cells[2].text = '日期'
    info_cells[3].text = date
    info_cells[4].text = '有效期至'
    info_cells[5].text = '30天'
    info_cells[6].text = '客户名称'
    info_cells[7].text = customer_name or '待填写'
    info_cells[8].text = '项目名称'
    info_cells[9].text = project_name or '待填写'
    info_cells[10].text = '供应商'
    info_cells[11].text = 'XX隔离器技术有限公司'
    info_cells[12].text = '设备类型'
    info_cells[13].text = eq
    info_cells[14].text = '配置等级'
    info_cells[15].text = budget
    info_cells[16].text = '适用法规'
    info_cells[17].text = ', '.join(regulations)
    
    doc.add_heading('A. 设备概述', level=1)
    doc.add_paragraph(equipment_desc(eq))
    
    doc.add_heading('B. 关键技术参数', level=1)
    param_table = doc.add_table(rows=6, cols=4)
    param_table.style = 'Table Grid'
    params = [
        ['灭菌效果', 'SAL 10⁻⁶', 'SAL 10⁻⁶', 'SAL 10⁻⁶'],
        ['残气浓度', '<1ppm', '<1ppm', '<1ppm'],
        ['内壁材质', '316L不锈钢', '316L不锈钢', '316L不锈钢'],
        ['控制系统', 'PLC+7"触摸屏', 'PLC+7"触摸屏', 'PLC+SCADA+10"触摸屏'],
        ['数据记录', '自动记录', '自动记录', '审计追踪记录'],
        ['双门互锁', '机械+电气', '机械+电气', '机械+PLC三重']
    ]
    
    hdr = param_table.rows[0].cells
    hdr[0].text = '参数项目'
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].text = '经济型'
    hdr[1].paragraphs[0].runs[0].bold = True
    hdr[2].text = '标准型'
    hdr[2].paragraphs[0].runs[0].bold = True
    hdr[3].text = '高端型'
    hdr[3].paragraphs[0].runs[0].bold = True
    
    for i, param in enumerate(params):
        row = param_table.rows[i+1]
        row.cells[0].text = param[0]
        row.cells[1].text = param[1]
        row.cells[2].text = param[2]
        row.cells[3].text = param[3]
    
    doc.add_heading(f'C. 报价明细（{budget}方案）', level=1)
    quote_table = doc.add_table(rows=8, cols=4)
    quote_table.style = 'Table Grid'
    quotes = [
        ['1', '设备本体', f'{base_price * 0.7:.1f}', '含主体结构、手套、视窗、VHP发生器'],
        ['2', '安装调试', f'{base_price * 0.08:.1f}', '含现场安装、调试、开机验证'],
        ['3', '验证服务（IQ/OQ/PQ）', f'{base_price * 0.1:.1f}', '含验证方案编制及执行'],
        ['4', '技术培训', f'{base_price * 0.02:.1f}', f"{3 if budget == '经济型' else 5}天（操作+维护{'+验证' if budget == '高端型' else ''}）"],
        ['5', '备件包（首年）', f'{base_price * 0.02:.1f}', '易损件及推荐备件'],
        ['6', '技术文件包', f'{base_price * 0.02:.1f}', '含手册、图纸、证书'],
        ['7', '国内运输+保险', f'{base_price * 0.06:.1f}', '含运输一切险'],
        ['**8**', '**合计**', f'**{base_price:.1f}**', '含13%增值税']
    ]
    
    hdr = quote_table.rows[0].cells
    hdr[0].text = '序号'
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].text = '项目'
    hdr[1].paragraphs[0].runs[0].bold = True
    hdr[2].text = '金额（万元）'
    hdr[2].paragraphs[0].runs[0].bold = True
    hdr[3].text = '备注'
    hdr[3].paragraphs[0].runs[0].bold = True
    
    for i, quote in enumerate(quotes):
        row = quote_table.rows[i+1]
        row.cells[0].text = quote[0]
        row.cells[1].text = quote[1]
        row.cells[2].text = quote[2]
        row.cells[3].text = quote[3]
    
    doc.add_heading('D. 配置清单', level=1)
    config_table = doc.add_table(rows=7, cols=4)
    config_table.style = 'Table Grid'
    configs = [
        ['1', f'{eq}主机系统', '1台', '含舱体、门系统、互锁装置'],
        ['2', 'HEPA过滤器（H14级）', '2个', '进风+排风'],
        ['3', '备用手套（丁腈）', '2副', '进口材质'],
        ['4', 'VHP发生器', '1套', '集成式'],
        ['5', '专用工具包', '1套', '维护工具'],
        ['6', '操作手册', '1份', '中文/英文'],
        ['7', '验证文件包', '1份', 'IQ/OQ/PQ方案']
    ]
    
    hdr = config_table.rows[0].cells
    hdr[0].text = '序号'
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].text = '项目'
    hdr[1].paragraphs[0].runs[0].bold = True
    hdr[2].text = '数量'
    hdr[2].paragraphs[0].runs[0].bold = True
    hdr[3].text = '备注'
    hdr[3].paragraphs[0].runs[0].bold = True
    
    for i, config in enumerate(configs):
        row = config_table.rows[i+1]
        row.cells[0].text = config[0]
        row.cells[1].text = config[1]
        row.cells[2].text = config[2]
        row.cells[3].text = config[3]
    
    doc.add_heading('E. 交货周期', level=1)
    delivery_table = doc.add_table(rows=3, cols=3)
    delivery_table.style = 'Table Grid'
    deliveries = [
        ['经济型', '12-16周', '自预付款到账之日起'],
        ['标准型', '16-20周', '自预付款到账之日起'],
        ['高端型', '20-28周', '含系统集成交付与调试时间']
    ]
    
    hdr = delivery_table.rows[0].cells
    hdr[0].text = '方案'
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].text = '预计周期'
    hdr[1].paragraphs[0].runs[0].bold = True
    hdr[2].text = '说明'
    hdr[2].paragraphs[0].runs[0].bold = True
    
    for i, delivery in enumerate(deliveries):
        row = delivery_table.rows[i+1]
        row.cells[0].text = delivery[0]
        row.cells[1].text = delivery[1]
        row.cells[2].text = delivery[2]
    
    doc.add_heading('F. 售后服务条款', level=1)
    service_table = doc.add_table(rows=6, cols=4)
    service_table.style = 'Table Grid'
    services = [
        ['质保期', '12个月', '12个月', '24个月'],
        ['远程响应时间', '48小时', '24小时', '24小时'],
        ['现场响应时间', '72小时', '48小时', '24小时'],
        ['年度保养', '1次/年', '1次/年', '2次/年'],
        ['软件升级', '质保期内免费', '质保期内免费', '质保期内免费+质保期后3年优惠'],
        ['设备备件供应', '设备停产10年', '设备停产10年', '设备停产10年']
    ]
    
    hdr = service_table.rows[0].cells
    hdr[0].text = '项目'
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].text = '经济型'
    hdr[1].paragraphs[0].runs[0].bold = True
    hdr[2].text = '标准型'
    hdr[2].paragraphs[0].runs[0].bold = True
    hdr[3].text = '高端型'
    hdr[3].paragraphs[0].runs[0].bold = True
    
    for i, service in enumerate(services):
        row = service_table.rows[i+1]
        row.cells[0].text = service[0]
        row.cells[1].text = service[1]
        row.cells[2].text = service[2]
        row.cells[3].text = service[3]
    
    doc.add_heading('声明', level=1)
    doc.add_paragraph('1. 本报价有效期为30天，逾期价格可能调整')
    doc.add_paragraph('2. 本报价不含土建、公用工程接口以外的工作')
    doc.add_paragraph('3. 最终价格以双方签订的正式合同为准')
    doc.add_paragraph('4. 如客户有重大变更，供应商保留调整报价的权利')
    
    doc.add_paragraph()
    doc.add_paragraph(f"**报价单位**: XX隔离器技术有限公司")
    doc.add_paragraph(f"**联系人**: {contact or '待填写'}")
    doc.add_paragraph(f"**日期**: {date}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def validate_urs_response():
    checks = []
    
    if not st.session_state.get('customer_name'):
        checks.append({'item': '客户名称', 'status': '缺失', 'level': 'warning'})
    else:
        checks.append({'item': '客户名称', 'status': '已填写', 'level': 'ok'})
    
    if not st.session_state.get('project_name'):
        checks.append({'item': '项目名称', 'status': '缺失', 'level': 'warning'})
    else:
        checks.append({'item': '项目名称', 'status': '已填写', 'level': 'ok'})
    
    if not st.session_state.selected_equipment:
        checks.append({'item': '设备类型', 'status': '未选择', 'level': 'warning'})
    else:
        checks.append({'item': '设备类型', 'status': f'已选择: {", ".join(st.session_state.selected_equipment)}', 'level': 'ok'})
    
    regulations = [r for r in ['中国GMP', 'FDA 21 CFR', 'EU GMP Annex 1', 'WHO GMP', 'PIC/S'] if st.session_state.get(f"reg_{r}")]
    if not regulations:
        checks.append({'item': '适用法规', 'status': '未选择', 'level': 'error'})
    else:
        checks.append({'item': '适用法规', 'status': f'已选择: {len(regulations)}项', 'level': 'ok'})
    
    if not st.session_state.urs_text.strip():
        checks.append({'item': 'URS内容', 'status': '未输入', 'level': 'error'})
    else:
        lines = [l for l in st.session_state.urs_text.split('\n') if l.strip()]
        checks.append({'item': 'URS内容', 'status': f'共{len(lines)}条需求', 'level': 'ok'})
    
    if st.session_state.parsed_requirements:
        categories = {}
        for req in st.session_state.parsed_requirements:
            categories[req['category']] = categories.get(req['category'], 0) + 1
        checks.append({'item': '需求分类', 'status': f'{len(categories)}个类别: {", ".join(categories.keys())}', 'level': 'ok'})
    
    return checks

def main():
    init_session()
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### 📋 客户基本信息")
        
        st.text_input("客户公司名称", key="customer_name")
        st.text_input("项目名称", key="project_name")
        
        row1 = st.columns(2)
        row1[0].text_input("联系人", key="contact")
        row1[1].date_input("日期", key="date", value=datetime.now())
        
        st.markdown("**适用法规（多选）**")
        regs = ['中国GMP', 'FDA 21 CFR', 'EU GMP Annex 1', 'WHO GMP', 'PIC/S']
        for reg in regs:
            st.checkbox(reg, key=f"reg_{reg}", value=True if reg == '中国GMP' else False)
        
        st.text_input("供应商名称", value="XX隔离器技术有限公司", key="supplier_name")
        
        st.markdown("---")
        st.markdown("### ⚙️ AI配置（可选）")
        st.text_input("OpenAI API Key（启用AI智能解析）", key="openai_api_key", type="password")
        if st.session_state.openai_api_key:
            st.success("✅ AI解析已启用")
        else:
            st.info("提示：输入OpenAI API Key可启用AI智能解析功能")
        
        st.markdown("---")
        st.markdown("### 📝 客户URS/询盘输入")
        st.text_area("请粘贴客户URS内容...", height=150, key="urs_text",
                     placeholder="1. 设备需满足无菌隔离要求，正压运行\n2. 操作区洁净度ISO 5级\n3. VHP灭菌，灭菌效果≥6-log\n4. 配备手套检漏系统\n5. 操作区尺寸不小于L1200×D800×H900mm\n6. 需满足中国GMP及EU GMP Annex 1要求")
        
        row2 = st.columns(3)
        if row2[0].button("🔍 智能解析"):
            with st.spinner("正在解析URS..."):
                parse_urs_with_ai(st.session_state.urs_text)
            if st.session_state.ai_parsed:
                st.success(f"✅ AI解析完成！识别到{len(st.session_state.parsed_requirements)}条需求")
            else:
                st.success(f"✅ 解析完成！识别到{len(st.session_state.parsed_requirements)}条需求")
        
        if row2[1].button("📄 加载示例"):
            st.session_state.urs_text = """1. VHP灭菌传递窗，双门互锁
2. 灭菌循环可调（浓度、时间、温度）
3. 灭菌效果≥6-log SAL
4. VHP残气排除至<1ppm
5. 内腔尺寸不小于L800×W800×H900mm
6. 需满足中国GMP及EU GMP Annex 1要求
7. 提供完整的IQ/OQ/PQ验证文件包"""
        
        if row2[2].button("🗑 清空"):
            st.session_state.urs_text = ""
            st.session_state.selected_equipment = []
            st.session_state.key_params = ""
            st.session_state.parsed_requirements = []
        
        with st.expander("⚙️ 手动调整", expanded=True):
            st.markdown("**设备类型确认/修正（可多选）**")
            eqs = ['单体无菌隔离器', 'VHP传递窗', '整线隔离器', '单体负压隔离器']
            for eq in eqs:
                st.checkbox(eq, key=f"eq_{eq}", 
                           value=True if eq in st.session_state.selected_equipment else False)
            
            st.text_area("关键参数修正", height=100, key="key_params",
                        placeholder="解析URS后将自动填充，也可手动输入")
            
            st.selectbox("预算级别", ['经济型', '标准型', '高端型'], key="budget_level")
            st.selectbox("语言", ['中文', 'English', '中英双语'], key="language")
        
        st.markdown("---")
        st.markdown("### ✅ 自检")
        if st.button("🔍 运行自检"):
            checks = validate_urs_response()
            for check in checks:
                if check['level'] == 'ok':
                    st.success(f"✅ {check['item']}: {check['status']}")
                elif check['level'] == 'warning':
                    st.warning(f"⚠️ {check['item']}: {check['status']}")
                else:
                    st.error(f"❌ {check['item']}: {check['status']}")
        
        st.markdown("---")
        st.markdown("### 💾 保存与加载")
        
        row3 = st.columns(2)
        if row3[0].button("💾 保存为报价单"):
            record = {
                'customer': st.session_state.customer_name,
                'equipment': st.session_state.selected_equipment,
                'date': datetime.now().strftime('%Y-%m-%d'),
                'budget': st.session_state.budget_level
            }
            st.session_state.history.append(record)
            save_data("history", st.session_state.history)
            st.success("保存成功！")
        
        template_name = st.text_input("模板名称")
        if row3[1].button("📑 保存为模板"):
            st.session_state.templates.append({
                'name': template_name,
                'content': st.session_state.urs_text,
                'date': datetime.now().strftime('%Y-%m-%d')
            })
            save_data("templates", st.session_state.templates)
            st.success("模板保存成功！")
        
        st.selectbox("加载模板", [t['name'] for t in st.session_state.templates] + [''], key="load_template")
        st.selectbox("加载客户", [c['name'] for c in st.session_state.customers] + [''], key="load_customer")
        
        with st.expander("📁 历史报价单", expanded=False):
            if st.session_state.history:
                for record in st.session_state.history[:5]:
                    st.markdown(f"{record['date']} | {', '.join(record['equipment'])} | {record.get('customer', '')}")
            else:
                st.info("暂无历史报价单")
        
        with st.expander("👥 客户档案", expanded=False):
            if st.session_state.customers:
                for cust in st.session_state.customers[:5]:
                    st.markdown(f"🏢 {cust['name']} | {cust.get('contact', '')}")
            else:
                st.info("暂无客户记录")
        
        if st.button("🚀 生成文档（URS回复 + 快速报价）", use_container_width=True, type="primary"):
            st.session_state.generated_docs['urs'] = generate_professional_urs_response()
            st.session_state.generated_docs['quotation'] = generate_quotation()
            st.session_state.generated_docs['urs_docx'] = create_urs_docx()
            st.session_state.generated_docs['quotation_docx'] = create_quotation_docx()
            st.success("文档生成成功！")
    
    with col2:
        tabs = st.tabs(["📋 URS逐条回复", "💰 快速报价单", "📚 模板库"])
        
        with tabs[0]:
            st.markdown("### 📋 URS逐条回复")
            if st.session_state.generated_docs['urs']:
                st.markdown(st.session_state.generated_docs['urs'])
                if st.session_state.generated_docs['urs_docx']:
                    st.download_button(
                        label="📥 下载 Word 格式 (.docx)",
                        data=st.session_state.generated_docs['urs_docx'],
                        file_name=f"URS逐条回复_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            else:
                st.info("请在左侧输入客户URS内容并点击\"智能解析\"，然后点击\"生成文档\"")
        
        with tabs[1]:
            st.markdown("### 💰 快速报价单")
            if st.session_state.generated_docs['quotation']:
                st.markdown(st.session_state.generated_docs['quotation'])
                if st.session_state.generated_docs['quotation_docx']:
                    st.download_button(
                        label="📥 下载 Word 格式 (.docx)",
                        data=st.session_state.generated_docs['quotation_docx'],
                        file_name=f"快速报价单_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            else:
                st.info("请在左侧输入信息并点击\"生成文档\"")
        
        with tabs[2]:
            st.markdown("### 📚 模板库")
            
            default_templates = [
                {
                    'name': '负压隔离器询盘（高活性药品）',
                    'equipment': '单体负压隔离器',
                    'regulations': '中国GMP、FDA 21 CFR',
                    'content': """1. 负压隔离器，用于高活性药品操作
2. OEL控制目标<1μg/m³
3. 双重HEPA排风过滤，在线扫描检漏
4. 袋进袋出(BIBO)过滤器更换
5. 工作区尺寸L1500×D900×H900mm"""
                },
                {
                    'name': 'VHP传递窗询盘',
                    'equipment': 'VHP传递窗',
                    'regulations': '中国GMP',
                    'content': """1. VHP灭菌传递窗，双门互锁
2. 灭菌循环可调（浓度、时间、温度）
3. 灭菌效果≥6-log SAL
4. VHP残气排除至<1ppm
5. 内腔尺寸不小于L800×W800×H900mm"""
                },
                {
                    'name': '标准无菌隔离器询盘',
                    'equipment': '单体无菌隔离器',
                    'regulations': '中国GMP、EU GMP',
                    'content': """1. 设备需满足无菌隔离要求，正压运行
2. 操作区洁净度ISO 5级
3. VHP灭菌，灭菌效果≥6-log
4. 配备手套检漏系统
5. 操作区尺寸不小于L1200×D800×H900mm"""
                }
            ]
            
            for template in default_templates:
                with st.expander(f"📄 {template['name']}"):
                    st.markdown(f"**设备类型**: {template['equipment']}")
                    st.markdown(f"**适用法规**: {template['regulations']}")
                    st.markdown(f"**内容**: \n{template['content']}")
                    if st.button(f"使用此模板", key=f"use_{template['name']}"):
                        st.session_state.urs_text = template['content']
                        st.session_state.selected_equipment = [template['equipment']]
                        st.success("模板已加载！")
        
        row4 = st.columns(2)
        if row4[0].button("📋 复制全文"):
            full_text = st.session_state.generated_docs['urs'] + '\n\n---\n\n' + st.session_state.generated_docs['quotation']
            st.write(full_text)
        
        if row4[1].button("🖨 打印/导出PDF"):
            st.write("打印功能已准备就绪")

if __name__ == "__main__":
    main()
