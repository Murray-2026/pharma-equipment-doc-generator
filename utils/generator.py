from templates.template_manager import TemplateManager
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


class DocumentGenerator:
    def __init__(self, equipment_type, regulation, custom_templates=None):
        self.template_manager = TemplateManager(equipment_type, regulation)
        self.equipment_type = equipment_type
        self.regulation = regulation
        self.custom_templates = custom_templates or {}
    
    def generate_urs_response(self, requirements=None):
        if self.custom_templates.get("urs_response"):
            return self._apply_custom_template("urs_response", requirements)
        
        if requirements and "urs_items" in requirements and requirements["urs_items"]:
            return self.generate_detailed_urs_response(requirements)
        return self.template_manager.get_urs_response_template(requirements)
    
    def generate_detailed_urs_response(self, requirements):
        content = f"# {self.equipment_type} - URS逐条回复\n\n"
        content += f"**文档编号**: URS-RESP-{datetime.now().strftime('%Y%m%d')}\n"
        content += f"**生成日期**: {datetime.now().strftime('%Y年%m月%d日')}\n"
        content += f"**适用法规**: {self.regulation}\n"
        content += f"**URS来源**: 客户提供的用户需求说明\n\n"
        content += "---\n\n"
        
        content += "## 1. 概述\n\n"
        content += "### 1.1 项目背景\n"
        content += f"本回复文档针对贵公司提出的{self.equipment_type}用户需求说明（URS）进行逐条响应，确保我们的方案完全满足贵公司的技术要求，并符合{self.regulation}相关法规要求。\n\n"
        content += "### 1.2 回复说明\n"
        content += "本文件采用逐条响应方式，对URS中的每个需求点进行详细说明，包括技术方案、合规性确认及相关备注信息。\n\n"
        content += "---\n\n"
        
        urs_items = requirements.get("urs_items", [])
        current_section = ""
        
        for item in urs_items:
            if item["type"] == "section":
                current_section = item["section_number"]
                content += f"## {item['section_number']} {item['title']}\n\n"
            elif item["type"] == "requirement":
                content += f"### {item['item_number']} 需求描述\n"
                content += f"{item['text']}\n\n"
                content += f"**我方响应**: {self.generate_response_for_requirement(item['text'])}\n\n"
                content += f"**合规状态**: ✅ 符合{self.regulation}要求\n\n"
                content += f"**验证方式**: 将在IQ/OQ/PQ验证中予以确认\n\n"
                content += "---\n\n"
        
        if not urs_items:
            content += "## 2. 基于提取的需求生成响应\n\n"
            content += self.template_manager.get_urs_response_template(requirements)
        
        content += "---\n\n"
        content += "## 附录\n\n"
        content += "### 术语说明\n"
        content += "- **URS**: 用户需求说明 (User Requirements Specification)\n"
        content += "- **IQ**: 安装确认 (Installation Qualification)\n"
        content += "- **OQ**: 运行确认 (Operational Qualification)\n"
        content += "- **PQ**: 性能确认 (Performance Qualification)\n\n"
        
        return content
    
    def _apply_custom_template(self, template_type, requirements=None):
        template = self.custom_templates.get(template_type, "")
        if not template:
            return ""
        
        placeholders = {
            "{equipment_type}": self.equipment_type,
            "{regulation}": self.regulation,
            "{date}": datetime.now().strftime("%Y年%m月%d日"),
            "{doc_number}": f"{template_type.upper()}-{datetime.now().strftime('%Y%m%d')}",
            "{year}": str(datetime.now().year),
        }
        
        if requirements:
            placeholders["{equipment_type_detected}"] = requirements.get("equipment_type", "未检测到")
            placeholders["{key_points_count}"] = str(len(requirements.get("key_points", [])))
        
        for placeholder, value in placeholders.items():
            template = template.replace(placeholder, value)
        
        return template
    
    def generate_response_for_requirement(self, requirement_text):
        responses = {
            "隔离器": "本公司提供的隔离器系统采用优质316L不锈钢材质，符合制药行业标准，具备完善的无菌隔离功能。",
            "无菌": "设备设计符合无菌工艺要求，配备高效HEPA H14过滤系统，确保A级洁净环境（ISO 5级）。",
            "VHP": "采用先进的VHP（汽化过氧化氢）灭菌技术，灭菌效率达到6-log，符合EN ISO 14644-7标准。",
            "传递窗": "传递窗设计符合GMP要求，具备互锁功能和VHP灭菌能力，确保洁净区完整性。",
            "整线": "整线隔离器系统可实现从物料进入到成品输出的全流程无菌操作，支持自动化集成。",
            "负压": "负压隔离器系统设计用于处理高活性或有害物料，具备袋进袋出（BIBO）过滤系统，确保操作人员安全。",
            "GMP": f"本设备完全符合{self.regulation}要求，可提供完整的验证文档包（DQ/IQ/OQ/PQ）。",
            "FDA": "设备设计符合FDA 21 CFR Part 11要求，支持电子记录和电子签名功能。",
            "EU GMP": "符合EU GMP Annex 1要求，适用于无菌药品生产，满足药品质量规范。",
            "PIC/S": "符合PIC/S GMP要求，支持国际多场地生产认证，确保全球合规。",
            "尺寸": "可根据客户需求定制设备尺寸，标准型号覆盖各种生产规模，确保满足生产工艺要求。",
            "参数": "设备关键参数可根据工艺需求进行调整和配置，提供灵活的解决方案。",
            "要求": "我方完全理解并满足客户提出的各项技术要求，确保方案的完整性和可行性。",
            "需求": "我方将根据客户需求提供定制化解决方案，确保满足项目的各项要求。",
            "验证": "提供完整的验证支持，包括IQ/OQ/PQ验证方案及报告模板，确保设备顺利通过验证。",
            "确认": "将通过正式的确认活动验证设备的性能和合规性，确保满足URS要求。",
        }
        
        for keyword, response in responses.items():
            if keyword in requirement_text:
                return response
        
        return "我方将根据客户需求提供相应的技术解决方案，确保满足项目要求并符合相关法规标准。"
    
    def generate_technical_spec(self):
        if self.custom_templates.get("technical_spec"):
            return self._apply_custom_template("technical_spec")
        
        return self.template_manager.get_technical_spec_template()
    
    def generate_quotation(self, config_type="standard"):
        if self.custom_templates.get("quotation"):
            return self._apply_custom_template("quotation")
        
        return self.template_manager.get_quotation_template(config_type)
    
    def generate_all_documents(self, requirements=None):
        documents = {
            "urs_response": self.generate_urs_response(requirements),
            "technical_spec": self.generate_technical_spec(),
            "quotation_basic": self.generate_quotation("basic"),
            "quotation_standard": self.generate_quotation("standard"),
            "quotation_premium": self.generate_quotation("premium")
        }
        return documents
    
    def save_markdown(self, content, filename, output_dir="outputs"):
        os.makedirs(output_dir, exist_ok=True)
        filepath = os.path.join(output_dir, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return filepath
    
    def create_word_document(self, content, filename, output_dir="outputs"):
        os.makedirs(output_dir, exist_ok=True)
        filepath = os.path.join(output_dir, filename)
        
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(10.5)
        
        lines = content.split("\n")
        in_table = False
        table_headers = []
        table_rows = []
        
        for line in lines:
            line = line.rstrip()
            if not line:
                doc.add_paragraph()
                continue
            
            if line.startswith("# "):
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in heading.runs:
                    run.font.size = Pt(16)
                    run.font.bold = True
            elif line.startswith("## "):
                heading = doc.add_heading(line[3:], level=2)
                for run in heading.runs:
                    run.font.size = Pt(14)
                    run.font.bold = True
            elif line.startswith("### "):
                heading = doc.add_heading(line[4:], level=3)
                for run in heading.runs:
                    run.font.size = Pt(12)
                    run.font.bold = True
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith("|") and "|" in line[1:]:
                in_table = True
                cells = [cell.strip() for cell in line.split("|") if cell.strip()]
                if not table_headers:
                    table_headers = cells
                else:
                    table_rows.append(cells)
            elif line.startswith("---") and in_table:
                continue
            elif line.startswith("- "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            else:
                if in_table and table_headers:
                    table = doc.add_table(rows=1, cols=len(table_headers))
                    table.style = "Table Grid"
                    
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(table_headers):
                        hdr_cells[i].text = header
                        for paragraph in hdr_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    for row_cells in table_rows:
                        row_cells = row_cells + [""] * (len(table_headers) - len(row_cells))
                        row = table.add_row().cells
                        for i, cell in enumerate(row_cells):
                            row[i].text = cell
                    
                    table_headers = []
                    table_rows = []
                    in_table = False
                
                doc.add_paragraph(line)
        
        if table_headers and table_rows:
            table = doc.add_table(rows=1, cols=len(table_headers))
            table.style = "Table Grid"
            
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(table_headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            for row_cells in table_rows:
                row_cells = row_cells + [""] * (len(table_headers) - len(row_cells))
                row = table.add_row().cells
                for i, cell in enumerate(row_cells):
                    row[i].text = cell
        
        doc.save(filepath)
        return filepath
