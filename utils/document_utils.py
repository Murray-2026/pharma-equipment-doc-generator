from docx import Document
from PyPDF2 import PdfReader
from io import BytesIO
from datetime import datetime
import re
import pandas as pd


def read_pdf(file):
    text = ""
    pdf_reader = PdfReader(file)
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text


def read_docx(file):
    doc = Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def read_excel(file):
    df = pd.read_excel(file)
    text = df.to_string(index=False)
    return text


def read_text(file):
    return file.read().decode("utf-8")


def create_docx_template(content, output_path):
    doc = Document()
    doc.add_heading("制药设备技术文档", 0)
    for section in content:
        if section.get("type") == "heading":
            doc.add_heading(section.get("text"), level=section.get("level", 1))
        elif section.get("type") == "paragraph":
            doc.add_paragraph(section.get("text"))
    doc.save(output_path)


def parse_urs_structure(text):
    urs_items = []
    current_section = ""
    current_section_title = ""
    
    lines = text.split("\n")
    line_number = 0
    
    for line in lines:
        line_number += 1
        line = line.strip()
        if not line:
            continue
        
        section_match = re.match(r'^(第?\s*(\d+)\s*章?|[\d.]+)\s*[、.．]?\s*(.+)$', line)
        if section_match:
            current_section = section_match.group(1)
            current_section_title = section_match.group(3)
            urs_items.append({
                "type": "section",
                "section_number": current_section,
                "title": current_section_title,
                "line_number": line_number
            })
            continue
        
        item_match = re.match(r'^([\d.]+)\s*[、.．]?\s*(.+)$', line)
        if item_match:
            item_number = item_match.group(1)
            item_text = item_match.group(2)
            urs_items.append({
                "type": "requirement",
                "section": current_section,
                "item_number": item_number,
                "text": item_text,
                "line_number": line_number
            })
        else:
            if urs_items and urs_items[-1]["type"] == "requirement":
                urs_items[-1]["text"] += " " + line
    
    return urs_items


def extract_requirements(text):
    requirements = {
        "equipment_type": None,
        "specifications": [],
        "compliance": [],
        "key_points": [],
        "urs_items": [],
        "raw_text": text
    }
    
    lines = text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        line_lower = line.lower()
        
        if "隔离器" in line or "isolator" in line_lower:
            if "无菌" in line_lower:
                requirements["equipment_type"] = "单体无菌隔离器"
            elif "VHP" in line or "传递窗" in line:
                requirements["equipment_type"] = "VHP传递窗"
            elif "整线" in line_lower or "line" in line_lower:
                requirements["equipment_type"] = "整线隔离器"
            elif "负压" in line_lower or "negative pressure" in line_lower:
                requirements["equipment_type"] = "单体负压隔离器"
        
        if re.search(r'(规格|尺寸|参数|spec|specification|dimension|size)', line_lower):
            requirements["specifications"].append(line)
        
        if re.search(r'(GMP|FDA|EU|PIC/S|ISO|合规|compliance|regulation)', line_lower):
            requirements["compliance"].append(line)
        
        if re.search(r'(要求|需求|必须|应|shall|should|must|need|require)', line_lower):
            requirements["key_points"].append(line)
    
    requirements["urs_items"] = parse_urs_structure(text)
    
    return requirements


def parse_uploaded_file(uploaded_file):
    file_type = uploaded_file.name.split(".")[-1].lower()
    
    if file_type == "pdf":
        text = read_pdf(uploaded_file)
    elif file_type in ["docx", "doc"]:
        text = read_docx(uploaded_file)
    elif file_type in ["xlsx", "xls"]:
        text = read_excel(uploaded_file)
    elif file_type == "txt":
        text = read_text(uploaded_file)
    else:
        text = ""
    
    requirements = extract_requirements(text)
    return text, requirements


def fill_template(uploaded_file, equipment_type, regulations, requirements=None):
    filename = uploaded_file.name
    file_type = filename.split(".")[-1].lower()
    
    # 准备占位符数据
    date_str = datetime.now().strftime("%Y年%m月%d日")
    year_str = str(datetime.now().year)
    regulation_str = ", ".join(regulations)
    doc_number = datetime.now().strftime("%Y%m%d")
    urs_doc_number = f"PHARMA-URS-{doc_number}"
    tech_doc_number = f"PHARMA-TECH-{doc_number}"
    quote_doc_number = f"PHARMA-QUOTE-{doc_number}"
    
    key_points_str = "\n".join(requirements.get("key_points", [])) if requirements else ""
    specs_str = "\n".join(requirements.get("specifications", [])) if requirements else ""
    
    placeholders = {
        "{设备类型}": equipment_type,
        "{法规标准}": regulation_str,
        "{生成日期}": date_str,
        "{年份}": year_str,
        "{URS文档编号}": urs_doc_number,
        "{技术文档编号}": tech_doc_number,
        "{报价文档编号}": quote_doc_number,
        "{关键需求点}": key_points_str,
        "{技术规格}": specs_str,
        "{文档编号}": doc_number,
        "{日期}": date_str
    }
    
    # 处理不同文件类型
    if file_type in ["docx", "doc"]:
        return fill_word_template(uploaded_file, placeholders)
    elif file_type in ["xlsx", "xls"]:
        return fill_excel_template(uploaded_file, placeholders)
    else:
        return fill_text_template(uploaded_file, placeholders)


def fill_word_template(uploaded_file, placeholders):
    from docx import Document
    from io import BytesIO
    
    doc = Document(uploaded_file)
    
    # 替换段落中的占位符
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
    
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in placeholders.items():
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)
    
    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    return {
        "file": output_buffer.getvalue(),
        "filename": f"filled_{uploaded_file.name}",
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }


def fill_excel_template(uploaded_file, placeholders):
    import pandas as pd
    from io import BytesIO
    
    df = pd.read_excel(uploaded_file)
    
    # 替换数据框中的占位符
    for col in df.columns:
        df[col] = df[col].astype(str)
        for placeholder, value in placeholders.items():
            df[col] = df[col].str.replace(placeholder, value)
    
    output_buffer = BytesIO()
    with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    
    output_buffer.seek(0)
    
    return {
        "file": output_buffer.getvalue(),
        "filename": f"filled_{uploaded_file.name}",
        "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    }


def fill_text_template(uploaded_file, placeholders):
    text = uploaded_file.read().decode("utf-8")
    
    for placeholder, value in placeholders.items():
        text = text.replace(placeholder, value)
    
    output_buffer = BytesIO()
    output_buffer.write(text.encode("utf-8"))
    output_buffer.seek(0)
    
    return {
        "file": output_buffer.getvalue(),
        "filename": f"filled_{uploaded_file.name}",
        "mime_type": "text/markdown" if uploaded_file.name.endswith(".md") else "text/plain"
    }
