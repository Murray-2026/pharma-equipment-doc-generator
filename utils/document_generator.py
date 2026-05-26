from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import tempfile
from datetime import datetime

class DocumentGenerator:
    def __init__(self, template_manager):
        self.template_manager = template_manager
    
    def generate(self, data):
        try:
            doc = Document()
            
            self.set_document_style(doc)
            
            self.add_title(doc, data)
            
            self.add_basic_info(doc, data)
            
            self.add_technical_specs(doc, data)
            
            self.add_features(doc, data)
            
            self.add_certifications(doc, data)
            
            self.add_footer(doc, data)
            
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            
            return temp_file.name
            
        except Exception as e:
            print(f"文档生成失败: {str(e)}")
            return None
    
    def set_document_style(self, doc):
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(11)
        
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    def add_title(self, doc, data):
        title = doc.add_heading('', level=0)
        run = title.add_run(f"{data['customer_name']}\n{data['equipment_type']}{data['doc_type']}")
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    def add_basic_info(self, doc, data):
        heading = doc.add_heading('项目基本信息', level=1)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        info_data = [
            ('项目名称', data['project_name']),
            ('客户名称', data['customer_name']),
            ('设备类型', data['equipment_type']),
            ('文档日期', data['date'])
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)
        
        doc.add_paragraph()
    
    def add_technical_specs(self, doc, data):
        heading = doc.add_heading('技术规格参数', level=1)
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        specs_data = [
            ('产能', data['capacity']),
            ('功率', data['power']),
            ('电压', data['voltage']),
            ('尺寸', data['dimensions']),
            ('材质', data['material'])
        ]
        
        for i, (label, value) in enumerate(specs_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value if value else '待定'
            
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)
        
        doc.add_paragraph()
    
    def add_features(self, doc, data):
        heading = doc.add_heading('主要功能特性', level=1)
        
        if data['features']:
            para = doc.add_paragraph(data['features'])
            para.paragraph_format.line_spacing = 1.5
        
        if data['application']:
            heading2 = doc.add_heading('典型应用场景', level=2)
            para = doc.add_paragraph(data['application'])
            para.paragraph_format.line_spacing = 1.5
        
        if data['advantages']:
            heading3 = doc.add_heading('竞争优势', level=2)
            para = doc.add_paragraph(data['advantages'])
            para.paragraph_format.line_spacing = 1.5
        
        doc.add_paragraph()
    
    def add_certifications(self, doc, data):
        if data['certification']:
            heading = doc.add_heading('认证要求', level=1)
            
            cert_text = '、'.join(data['certification'])
            para = doc.add_paragraph(f"本设备满足以下认证要求：{cert_text}")
            para.paragraph_format.line_spacing = 1.5
            
            doc.add_paragraph()
    
    def add_footer(self, doc, data):
        doc.add_paragraph()
        
        footer = doc.add_paragraph()
        footer.add_run('─' * 50)
        
        footer_text = doc.add_paragraph()
        footer_text.add_run(f"文档生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        footer_text2 = doc.add_paragraph()
        footer_text2.add_run("制药设备售前文档生成器")
        footer_text2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
